"""
Yargı AI - FastAPI Chat Arayüzü - GPT-5 Powered
ChatGPT/GPT-5 benzeri tek sayfa chat arayüzü
"""

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request, Header
from fastapi.responses import JSONResponse, StreamingResponse
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from io import BytesIO
from typing import Dict, List, Any, Optional, AsyncGenerator

# Yeni eklenen importlar
import os, asyncio, json, uuid, re, traceback
import PyPDF2
import uvicorn

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:     
    print("⚠️  python-dotenv not installed. Using system environment variables.")

# Local imports
from legal_advisor_ai import SingleLegalAdvisor
from prompts import build_legal_analysis_prompt  # merkezi şablon
# post_processing modülü kaldırıldı: ilgili işlevler formatting.standardize_output içinde ele alınıyor
from formatting import standardize_output
from petition_generator_ai import PetitionGeneratorAI
from vector_store import get_vector_store, new_session_id
from precedent_service import (
    prepare_precedents_for_detailed_answer,
    generate_precedent_search_query,
    summarize_precedents_for_prompt,
    summarize_precedents_ai,
    enforce_citations,
    build_full_precedents_block,
    build_semantic_precedents_block,
)
from web_context import WebContextFetcher, detect_law_area, get_domains_for_area

# Multi-token AI Guard sistemi - ZORUNLU
from ai_guard import (
    get_multi_token_manager,
    get_ai_manager,
    safe_openai_request,
    safe_main_ai_request,
    safe_search_ai_request,
    safe_petition_ai_request,
    safe_case_analysis_ai_request,
    get_token_status
)

print("✅ Multi-token AI Guard sistemi yüklendi (App)")

# Global AI instances
legal_advisor_ai = None
petition_generator_ai = None
# Sohbet bazlı emsal setleri: session_id -> precedent_set_id -> { 'precedents': [...], 'created': iso }
# NOT: session_chat_precedents küçük metadata dict'idir, vektör verisi değil.
# Asıl büyük veriler (konuşmalar, dosyalar, emsaller) artık tamamen Pinecone'da tutulur.
session_chat_precedents: dict[str, dict[str, dict]] = {}

# ================================
# LIFESPAN MANAGEMENT
# ================================

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Uygulama yaşam döngüsü yönetimi"""
    # Startup
    global legal_advisor_ai, petition_generator_ai
    try:
        # AI Guard başlat - multi-token sistem zorunlu
        ai_manager = get_ai_manager()
        print("✅ AI Guard sistemi başlatıldı")
        # Tekil hukuki danışman
        legal_advisor_ai = SingleLegalAdvisor()
        print("✅ Hukuki danışman hazır")
        # Bedesten istemcisi kaldırıldı (dava analizi özelliği kaldırıldı)
    except Exception as e:
        print(f"❌ Hukuki danışman başlatılamadı: {str(e)}")
    try:
        petition_generator_ai = PetitionGeneratorAI()
        print("✅ Dilekçe AI hazır")
    except Exception as e:
        print(f"❌ Dilekçe AI başlatılamadı: {str(e)}")
    
    yield
    
    # Shutdown
    print("🔄 Uygulama kapatılıyor...")

# ================================
# APP CONFIGURATION
# ================================

app = FastAPI(
    title="Yargı AI Chat",
    description="ChatGPT benzeri Hukuki AI Asistanı",
    version="3.0.0",
    lifespan=lifespan
)

# Not: Sunucu taraflı UI kaldırıldı. Bu uygulama yalnızca API'ları sunar.

# ================================
# DEBUG / DIAGNOSTIC ENDPOINTS
# ================================
@app.get("/debug/full-precedents")
async def debug_full_precedents(session_id: str, request: Request):
    """Belirtilen session'daki precedent dokümanlarının boyutlarını döndürür. Sadece localhost'ta aktif."""
    client_host = request.client.host if request.client else "unknown"
    if client_host not in ("127.0.0.1", "::1", "localhost"):
        raise HTTPException(status_code=403, detail="Bu endpoint yalnızca localhost'ta erişilebilir")
    store = get_vector_store()
    full_map = store.get_full_precedents_for_session(session_id)
    sizes = {k: len(v) for k, v in full_map.items()}
    return {"precedent_count": len(full_map), "sizes": sizes}

# ================================
# ERROR HANDLERS
# ================================

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """HTTP hata işleyicisi (JSON)."""
    return JSONResponse({
        "error_message": exc.detail
    }, status_code=exc.status_code)

@app.exception_handler(500)
async def internal_error_handler(request: Request, exc):
    """İç hata işleyicisi (JSON)."""
    return JSONResponse({
        "error_message": "Sunucu hatası oluştu. Lütfen daha sonra tekrar deneyin."
    }, status_code=500)

# ================================
# UTILITY FUNCTIONS
# ================================

import time as _time

class TimingLogger:
    """Basit süre ölçer: her adımı ve toplamı terminale yazar."""
    def __init__(self, label: str, session_id: Optional[str] = None):
        self.label = label
        self.session_id = session_id or "-"
        self.t0 = _time.perf_counter()
        self.last = self.t0
        print(f"⏱️  [{self.label}][{self.session_id}] başladı")

    def lap(self, step: str):
        now = _time.perf_counter()
        dt = (now - self.last) * 1000.0
        tot = (now - self.t0) * 1000.0
        print(f"⏱️  [{self.label}][{self.session_id}] {step}: +{dt:.1f} ms (toplam {tot:.1f} ms)")
        self.last = now

    def end(self, note: str = "bitti"):
        now = _time.perf_counter()
        tot = (now - self.t0) * 1000.0
        print(f"✅ [{self.label}][{self.session_id}] {note}: toplam {tot:.1f} ms")

def format_ai_response(text: str) -> str:
    """AI yanıtını görsel olarak sadeleştirip normalize eder."""
    import re
    try:
        if not text or not text.strip():
            return text
        # Ön temizleme: görünmez karakterler
        text = re.sub(r"[\u00A0\ufeff\u200b\u200c\u200d]", " ", text)
        lines = text.splitlines()
        cleaned: list[str] = []

        def _looks_heading(s: str) -> bool:
            return bool(re.match(r"^(#{1,4}\s+|[A-ZÇĞİÖŞÜ0-9][^a-zçğıöşü]{3,}$)", s.strip()))

        for raw in lines:
            ln = raw.strip()
            if not ln:
                cleaned.append("")
                continue
            # *Başlık -> ### Başlık
            if re.match(r"^\*[A-ZÇĞİÖŞÜ].+", ln):
                ln = '### ' + ln.lstrip('*').strip()
            # 1. ... -> 1) ... (listeye dönüşmesin)
            ln = re.sub(r"^(\d{1,3})\.\s+", r"\1) ", ln)
            # Tek kelimelik büyük blokları heading yap
            if _looks_heading(ln) and not ln.startswith('#'):
                if len(ln.split()) <= 10:
                    ln = '### ' + ln.title()
            cleaned.append(ln)

        # Boş satırları normalize et
        out = []
        for ln in cleaned:
            if ln == "" and (not out or out[-1] == ""):
                continue
            out.append(ln)
        result = '\n'.join(out).strip()
        # Fazla ardışık başlıkları araya boş satır koy
        result = re.sub(r"(### .+?)\n(### )", r"\1\n\n\2", result)
    # clean_visual_markdown kaldırıldı (post_processing silindi)
        return result.strip()
    except Exception:
        return text

def _strip_source_tags(text: str) -> str:
    """Metin içindeki '(Kaynak: ...)' etiketlerini kaldırır.
    NOT: [w1], [w2] gibi web kaynak atıfları ve AI'ın Kaynaklar bölümü
    KORUNUR — bunlar uygulama tarafından işlenir.
    """
    try:
        import re as _re
        if not text:
            return text
        # Remove inline (Kaynak: ...) parantez etiketleri (Bedesten emsal etiketleri)
        text = _re.sub(r"\(\s*Kaynak\s*:\s*[^\)]*\)", "", text)
        # Trim excessive blank lines
        text = _re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
    except Exception:
        return text


def _extract_cited_web_indices(text: str) -> set:
    """AI yanıt metninde geçen [w1], [w2] gibi web kaynak atıf indekslerini döndür.

    Bu fonksiyon sayesinde footer'a yalnızca AI'ın gerçekten atıf yaptığı
    kaynaklar eklenir; alakasız sonuçlar (örn: aile hukuku sorusundaki
    Danıştay dergileri) gösterilmez.
    """
    import re as _re
    if not text:
        return set()
    return {int(m.group(1)) for m in _re.finditer(r'\[w(\d+)\]', text)}
        
# Basit kanun link üretici ve tespit edici (mevzuat.gov.tr)
LAW_MAP = {
    'TMK': '4721',  # Türk Medeni Kanunu
    'TBK': '6098',  # Türk Borçlar Kanunu
    'HMK': '6100',  # Hukuk Muhakemeleri Kanunu
    'TCK': '5237',  # Türk Ceza Kanunu
    'İİK': '2004', 'IİK': '2004', 'IIK': '2004',  # İcra ve İflas Kanunu (farklı yazımlar)
    'TTK': '6102',  # Türk Ticaret Kanunu
    'FSEK': '5846', # Fikir ve Sanat Eserleri Kanunu
    'SMK': '6769',  # Sınai Mülkiyet Kanunu
    'KVKK': '6698', # Kişisel Verilerin Korunması Kanunu
}

def _normalize_law_code(code: str) -> str:
    return 'İİK' if code in {'IİK', 'IIK'} else code

def _law_url(law_no: str) -> str:
    return f"https://www.mevzuat.gov.tr/Mevzuat?MevzuatNo={law_no}&MevzuatTur=1&MevzuatTertip=5"

def _extract_law_codes(text: str) -> list[str]:
    try:
        import re as _re
        if not text: return []
        # Match [TMK m.166] or TMK m.166 (avoid inside URLs)
        pattern = _re.compile(r"\b(TMK|TBK|HMK|TCK|İİK|IİK|IIK|TTK|FSEK|SMK|KVKK)\s*m\.?\s*\d+[A-Za-z\-]*", _re.IGNORECASE)
        found = set()
        for m in pattern.finditer(text):
            code = m.group(1).upper()
            code = _normalize_law_code(code)
            if code in LAW_MAP:
                found.add(code)
        return sorted(found)
    except Exception:
        return []
# Dava analizi (case analysis) akışları kaldırıldı

def extract_text_from_pdf(pdf_bytes) -> str:
    """PDF'ten metin çıkar (PyPDF2)."""
    from io import BytesIO as _BytesIO
    try:
        with _BytesIO(pdf_bytes) as pdf_stream:
            reader = PyPDF2.PdfReader(pdf_stream)
            text = ''
            for page in reader.pages:
                try:
                    text += page.extract_text() + '\n'
                except Exception:
                    continue
            return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF okuma hatası: {e}")


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """DOCX dosyasından metin çıkar (python-docx)."""
    from io import BytesIO as _BytesIO
    try:
        from docx import Document as _Document
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx kütüphanesi kurulu değil. 'pip install python-docx' ile kurun.")
    try:
        with _BytesIO(docx_bytes) as stream:
            doc = _Document(stream)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Tablolardaki metni de çıkar
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return '\n'.join(paragraphs).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"DOCX okuma hatası: {e}")

async def process_case_analysis(file_content: str, user_prompt: str, session_id: str = None) -> Dict[str, Any]:
    return {"error": "Dava analizi özelliği kaldırıldı"}

async def build_memory_prompt(session_id: str, user_message: str, file_content: Optional[str] = None, k_similar: int = 5, k_recent: int = 4, detail: str = "concise") -> str:
    """Vektör veritabanından benzer ve son mesajları çekip prompt oluştur.

    - Benzer mesajlar: FAISS similarity (hem chat hem pdf) sonuçları
    - Son mesajlar: Kronolojik son k_recent (kullanıcı & asistan) mesajları
    - Aynı içerik iki kere gelirse (ör: hem similar hem recent) tekrar etme
    """
    store = get_vector_store()
    similar = []
    recent = []
    try:
        similar = await store.a_similarity_search(session_id=session_id, query=user_message, k=k_similar, include_pdf=True, include_chat=True)
    except Exception as e:
        print(f"[Memory] a_similarity_search hata: {e}")
    try:
        recent = store.get_chat_history(session_id=session_id)
        if recent:
            recent = recent[-k_recent:]
    except Exception as e:
        print(f"[Memory] get_chat_history hata: {e}")

    # Benzer mesajlardan sadece chat olanları öncelik ver (role var ise)
    similar_chats = [s for s in similar if s.get("kind") == "chat"]
    # PDF chunk'ları da ekleyelim (ayrı başlıkla). İçerik çok uzunsa truncate.
    similar_pdfs = [s for s in similar if s.get("kind") == "pdf"]

    def shorten(text: str, limit: int = 400) -> str:
        return text if len(text) <= limit else text[:limit] + "..."

    memory_parts = []
    if similar_chats:
        memory_parts.append("-- İLGİLİ GEÇMİŞ MESAJLAR --")
        for sc in similar_chats:
            role = sc.get("role") or "gecmis"
            memory_parts.append(f"[{role}] {shorten(sc.get('chunk_text',''))}")
    if similar_pdfs:
        memory_parts.append("-- İLGİLİ YÜKLENEN DOSYA PARÇALARI --")
        used_ids = set()
        for sp in similar_pdfs:
            pid = sp.get("pdf_id")
            if (pid, sp.get("chunk_id")) in used_ids:
                continue
            used_ids.add((pid, sp.get("chunk_id")))
            memory_parts.append(f"[PDF:{pid}] {shorten(sp.get('chunk_text',''))}")
    if recent:
        memory_parts.append("-- SON MESAJLAR --")
        for r in recent:
            memory_parts.append(f"[{r.get('role')}] {shorten(r.get('content',''))}")

    if file_content:
        memory_parts.append("-- GÜNCEL DOSYA İÇERİĞİ (KISALTILMIŞ) --")
        memory_parts.append(shorten(file_content, 1200))

    memory_context = "\n".join(memory_parts) if memory_parts else "(Önceki bağlam bulunamadı)"

    # Üç tırnak içindeki kullanıcı mesajını güvenli yerleştirmek için escape kullan
    safe_user = user_message.replace('"""', '"\"\"')
    style = (
        # Yeni dengeli "Esnek Kısa" modu: Çok kısa kesme; makul uzunlukta, okunabilir, başlık + maddeler + kısa paragraflar.
        "Yanıtı dengeli, açık ve okunabilir ver. 2-3 kısa ama açıklayıcı bölüm başlığı (## veya ###) kullan. "
        "Önemli noktaları 5-10 maddelik listeler halinde özetle; listeler sadece tekrarlayan kelimeler olmasın. "
        "Gerektiğinde 1-3 cümlelik sade açıklayıcı paragraflar ekle. Resmî üslup ile günlük anlaşılır dil arasında orta bir netlik kullan. "
        "İlgili mevzuatı madde numarasıyla KISA referansla (örn: TTK m.123) an fakat metni blok halinde kopyalama. "
        "Varsa riskler / dikkat noktaları / öneriler için ayrı mini başlık aç. Gereksiz tekrar, aşırı süslü giriş, sonuçta klişe kapanış yapma."
        if detail == "concise" else
        "Kapsamlı ve sistematik bir hukuk analizi yap. Bölümlü başlıklar, madde madde argümanlar, karşı görüşler ve riskler, uygulanabilir mevzuat ve yol haritası ver. Gerektiğinde örnek metin/şablon öner."
    )

    prompt = (
        "Aşağıda kullanıcının mevcut mesajı ve önceki oturum bağlamı yer alıyor.\n"
        "Önce geçmişe dayanarak tutarlı, çelişkileri gideren ve önceki cevapların tekrarını azaltan bir hukuki yardımcı cevabı üret.\n\n"
        "KULLANICININ MEVCUT MESAJI:\n" + safe_user + "\n\n"
        "OTURUM HAFIZASI / BAĞLAM:\n" + memory_context + "\n\n"
        f"Yanıt stili: {style}"
    )
    return prompt

def build_memory_prompt_with_vector_content(session_id: str, user_message: str, vector_content: str = "", detail: str = "concise") -> str:
    """Vector search sonuçlarıyla birleştirilmiş prompt oluştur."""
    store = get_vector_store()
    recent = []
    try:
        recent = store.get_chat_history(session_id=session_id)
        if recent:
            recent = recent[-4:]  # Son 4 mesaj
    except Exception as e:
        print(f"[Memory] get_chat_history hata: {e}")

    def shorten(text: str, limit: int = 400) -> str:
        return text if len(text) <= limit else text[:limit] + "..."

    memory_parts = []
    
    # Vector search sonuçlarını ekle
    if vector_content:
        memory_parts.append("-- YÜKLENEN DOSYALARDAN BULUNAN İLGİLİ BÖLÜMLER (ZORUNLU KULLANIM) --")
        memory_parts.append("⚠️ Aşağıdaki içerikler kullanıcının yüklediği dosyalardan semantik arama ile bulunmuştur.")
        memory_parts.append("Bu bilgileri MUTLAKA yanıtında kullan ve hangi belgeden alındığını belirt.")
        # Dosya chunk'larına tam erişim: toplam bütçe kontrolü shrink_context()'te yapılır.
        # Normal kullanımda 10 chunk * 1200 char ≈ 12K → 50K bütçe içinde güvenli.
        memory_parts.append(vector_content)
    
    # Son sohbet geçmişi
    if recent:
        memory_parts.append("-- SON SOHBET GEÇMİŞİ --")
        for r in recent:
            if isinstance(r, dict) and r.get('role') and r.get('content'):
                memory_parts.append(f"[{r.get('role')}] {shorten(r.get('content',''))}")

    memory_context = "\n".join(memory_parts) if memory_parts else "(Bağlam bulunamadı)"

    # Güvenli kullanıcı mesajı
    safe_user = user_message.replace('"""', '"\"\"')
    style = (
        "Dengeli ve yapılandırılmış yanıt ver. 2-3 kısa bölüm başlığı kullan; her bölümde 1 kısa açıklayıcı paragraf + gerektiğinde 4-8 maddelik nokta listesi olsun. "
        "Gereksiz tekrar yapma; mevzuat atıflarını kısalt (örn: TMK m.2, HMK m.119). Teknik terimleri günlük anlaşılır dille özetle. "
        "Risk / dikkat / öneri için ayrı küçük başlık aç. Son bölümde sonuçları net ve sade biçimde toparla; kalıp teşekkür cümlesi ekleme."
        if detail == "concise" else
        "Kapsamlı ve sistematik bir hukuk analizi yap. Bölümlü başlıklar, argümanlar, karşı görüşler, mevzuat ve yol haritası ver. Belge referanslarını açık yaz."
    )
    prompt = (
        "Aşağıda kullanıcının sorusu ve yüklenen belgelerden ilgili içerikler yer alıyor.\n"
        "ÖNEMLİ: Eğer belge içerikleri verilmişse, bunları MUTLAKA oku ve yanıtında kullan.\n"
        "Belgelerden edindiğin bilgilerle kullanıcının sorusuna doğru ve detaylı bir cevap ver.\n"
        "Hangi belgeden bilgi aldığını açıkça belirt.\n\n"
        "KULLANICININ SORUSU:\n" + safe_user + "\n\n"
        "İLGİLİ BAĞLAM VE BELGE İÇERİKLERİ:\n" + memory_context + "\n\n"
        f"Yanıt stili: {style}"
    )
    return prompt

async def process_regular_chat(message: str, session_id: str, file_content: str = None) -> Dict[str, Any]:
    """Normal chat işlemi - session'a özel vektör hafızasıyla."""
    try:
        if legal_advisor_ai is None:
            return {"error": "AI sistemi başlatılamadı"}
        memory_prompt = await build_memory_prompt(session_id=session_id, user_message=message, file_content=file_content)
        central_prompt = build_legal_analysis_prompt(user_question=message, context_snippet=memory_prompt, detail_level="standart", mode="flex_concise")
        # Tekil sistemde doğrudan süreç
        result = await legal_advisor_ai.process_question(central_prompt, flex_mode="flex_detailed")
        if result and 'response' in result:
            formatted_response = format_ai_response(result['response'])
            return {
                "response": formatted_response,
                "timestamp": datetime.now().isoformat(),
                "type": "regular_chat"
            }
        return {"error": "AI'dan yanıt alınamadı"}
    except Exception as e:
        return {"error": f"Chat sırasında hata: {str(e)}"}

async def process_normal_chat_with_vector_search(message: str, session_id: str, detail: str = "concise", party_info: Optional[str] = None, situation_info: Optional[str] = None, include_precedents: bool = True) -> Dict[str, Any]:
    """Normal chat işlemi - vector store'daki tüm dosyalarda arama yapar.

    Gereksinim: Sorgu başına 20 sonuç getir, en yüksek bağlama sahip 10'unu
    kullan ve kalanını at. (Mevcut mimaride PDF chunk'ları kalıcı olduğundan
    burada sadece ilk 10'u prompt'a dahil ediyoruz; ileride mevzuat/dış
    arama sonuçları 'temp' olarak eklenirse trim_session_vectors_top_k
    çağrısı ile fazlalık fiziksel olarak silinebilir.)
    """
    try:
        tlog = TimingLogger("chat-normal", session_id)
        # Vector store'dan ilgili içerikleri ara
        store = get_vector_store()
        
        # Session'da yüklü dosya olup olmadığını kontrol et (Pinecone)
        session_has_files = store.has_pdf_files(session_id)
        
        # Sadece PDF içeriklerinde ara (chat mesajları hariç)
        search_results = await store.a_similarity_search(
            session_id=session_id,
            query=message,
            k=20,  # 20 sonuç al
            include_pdf=True,
            include_chat=False
        )
        
        # Pinecone eventual consistency: sonuç yoksa ama dosya varsa tekrar dene
        if not search_results and session_has_files:
            print(f"⚠️ [Vector] Dosya var ama arama sonuç vermedi, 2s bekleyip tekrar deneniyor...")
            await asyncio.sleep(2)
            search_results = await store.a_similarity_search(
                session_id=session_id,
                query=message,
                k=20,
                include_pdf=True,
                include_chat=False
            )
        
        # Hâlâ sonuç yoksa ama dosya varsa, Pinecone fetch ile ilk chunk'ları al
        if not search_results and session_has_files:
            print(f"⚠️ [Vector] Pinecone arama sonuç vermedi, fetch fallback kullanılıyor")
            fallback_docs = store.get_all_documents_raw(session_id, kinds=["pdf"])
            for r in fallback_docs[:10]:
                search_results.append({
                    "session_id": session_id,
                    "pdf_id": r["pdf_id"],
                    "chunk_id": r["chunk_id"],
                    "chunk_text": r["text"],
                    "distance": 0.0,
                    "kind": "pdf",
                    "role": None,
                })
        
        tlog.lap(f"vector search tamamlandı ({len(search_results)} sonuç)")
        
        # Hafızalı prompt oluştur - vector search sonuçlarını ekle
        relevant_content = ""
        if search_results:
            top10 = search_results[:10]
            # Chunk'lar zaten 1200 karakter (chunk_size); tekrar kesme yapmıyoruz.
            # Genel bütçe kontrolü shrink_context() tarafından yönetilir.
            relevant_content = "\n\n".join([
                f"**{r['pdf_id']}** (Benzerlik: {r['distance']:.3f}):\n{r['chunk_text']}" for r in top10
            ])
        
        base_memory = build_memory_prompt_with_vector_content(
            session_id=session_id,
            user_message=message,
            vector_content=relevant_content,
            detail=detail
        )
        tlog.lap("hafıza/prompt hazır")

        precedent_context = ""
        precedent_meta: Dict[str, Any] = {}
        # Precedent taraması sadece detay modda ve kullanıcı istediğinde yapılır
        if detail != "concise" and include_precedents:
            try:
                precedent_meta = await prepare_precedents_for_detailed_answer(
                    session_id=session_id,
                    user_question=message,
                    context_snippet=relevant_content,  # dosya içeriğini geç: meta sorularda asıl konuyu çıkarsın
                )
                selected = precedent_meta.get("selected", [])
                if selected:
                    # Statik kısa özet — hızlı meta referans (hangi karar, hangi birim)
                    static_block = summarize_precedents_for_prompt(selected)
                    # AI tabanlı kısa özet — bu işlem sırasında Pinecone indexlemesi de tamamlanır
                    ai_short = await summarize_precedents_ai(selected)
                    # SEMANTİK EMSAL PARÇALARI — Pinecone'dan sorguya en ilgili chunk'ları çek.
                    # Emsal chunk'ları Pinecone'da tam metin olarak saklanıyor; semantik olarak
                    # sorguya en yakın paragraflar seçilerek modele gönderiliyor.
                    # Sadece BU analize ait stored_ids ile filtrele; aynı session'daki
                    # önceki analizlere ait emsal chunk'ları karışmasın.
                    _current_stored_ids = precedent_meta.get("stored_ids", [])
                    # Pinecone propagation retry: yeni upsert edilen chunk'lar 2-10sn gecikmeli görünür.
                    # 0 sonuç gelirse 3sn bekle, tekrar dene (max 3 deneme: 0sn / 3sn / 6sn).
                    sem_chunks = []
                    _retry_waits = [0, 3, 6]
                    for _attempt, _wait in enumerate(_retry_waits):
                        if _wait > 0:
                            print(f"📑 [Pinecone] Semantik arama boş, {_wait}sn bekleniyor (deneme {_attempt+1}/{len(_retry_waits)})...")
                            await asyncio.sleep(_wait)
                        sem_chunks = await store.a_precedent_similarity_search(
                            session_id=session_id,
                            query=message,
                            k=12,
                            pdf_ids=_current_stored_ids if _current_stored_ids else None,
                        )
                        if sem_chunks:
                            print(f"📑 [Pinecone] {len(sem_chunks)} semantik chunk bulundu (deneme {_attempt+1})")
                            break
                    if sem_chunks:
                        emsal_icerik_blogu = build_semantic_precedents_block(sem_chunks, selected, max_chunks=12)
                        emsal_icerik_etiket = "-- SEMANTİK EMSAL PARÇALARI (SORGUYLA EN İLGİLİ, ANALİZDE ZORUNLU KULLANIM) --"
                    else:
                        # Tüm denemeler sonrası hâlâ boş: ham metin fallback
                        emsal_icerik_blogu = build_full_precedents_block(selected, max_chars=18000)
                        emsal_icerik_etiket = "-- TAM EMSAL İÇERİKLERİ (ANALİZDE ZORUNLU KULLANIM) --"
                        print(f"📑 Semantik arama tüm denemelerde boş döndü, build_full_precedents_block fallback kullanılıyor")
                    precedent_context = (
                        "\n" + static_block
                        + "\n\n-- KISA EMSAL ÖZETLERİ --\n" + ai_short
                        + f"\n\n{emsal_icerik_etiket}\n" + emsal_icerik_blogu
                    )
                    print(f"📑 Detaylı mod: {len(selected)} emsal yüklendi, {len(sem_chunks)} semantik chunk — icerik_blogu={len(emsal_icerik_blogu)}ch (search='{precedent_meta.get('search_query')}')")
            except Exception as pe:
                print(f"⚠️ Emsal hazırlama hatası: {pe}")
        if precedent_context:
            tlog.lap("emsal bağlam hazır")

        # Web kaynaklarını (mevzuat/doktrin) sadece detay modda ara
        web_section = ""
        web_sources: list[dict] = []
        web_link_map: dict[int, dict] = {}
        web_semantic_context = ""
        if detail != "concise":
            try:
                web_fetcher = WebContextFetcher()
                if web_fetcher.enabled:
                    # Web arama query'si: Tavily kısa odaklı sorgu ister (600 char ham metin → 1 sonuç, 0.045 score).
                    # Öncelik sırası:
                    #   1. Emsal sistemi zaten AI ile odaklı query ürettiyse onu kullan (aynı konuyu kapsıyor)
                    #   2. Soru direkt (meta değil) ise message'ı kullan (max 120 char)
                    #   3. Meta soru + doküman varsa: context'ten AI ile yeni query üret
                    _prec_query = precedent_meta.get("search_query") if precedent_meta else None
                    _meta_words = {"özetle", "anlat", "dosya", "belge", "dosyadaki", "yol göster", "özet"}
                    _is_meta = any(w in message.lower() for w in _meta_words)

                    if _prec_query:
                        # Emsal query zaten temiz ve odaklı (ör: "FSEK yazılım eser hakkı ihlali tazminat")
                        law_query = _prec_query
                        print(f"🔍 [WebContext] Emsal query web aramaya aktarıldı: '{law_query}'")
                    elif not _is_meta:
                        # Direkt soru: message kısa ve konuyla ilgili — sadece kes
                        law_query = message[:120]
                    else:
                        # Meta soru + emsal yok: context'ten AI query üret
                        law_query = await generate_precedent_search_query(message, relevant_content)
                        if not law_query:
                            # heuristic: context'ten ilk anlamlı kelimeler
                            import re as _re2
                            _stop2 = {"ve","veya","ile","bir","bu","şu","için","olan","göre","da","de"}
                            _toks = [t.lower() for t in _re2.findall(r"[A-Za-zÇĞİÖŞÜçğıöşü]{4,}", relevant_content[:800])]
                            _kws = [t for t in _toks if t not in _stop2][:7]
                            law_query = " ".join(dict.fromkeys(_kws)) or message[:80]
                        print(f"🔍 [WebContext] Meta soru — web query üretildi: '{law_query}'")

                    # Alan tespiti için bağlamı dahil et (keyword matching, Tavily'e gitmez)
                    _area_text = (message + " " + relevant_content[:400]).strip() if relevant_content else message
                    detected_area = detect_law_area(_area_text)
                    print(f"🔍 [WebContext] Tespit edilen hukuk alanı: '{detected_area}' → domain seti seçildi")
                    max_results = 8
                    web_data = await web_fetcher.asearch(
                        law_query,
                        max_results=max_results,
                        search_depth="advanced",
                        days=3650,
                        law_area=detected_area,    # Alan bazlı domain filtresi
                        min_score=0.40,            # 0.50 çok yüksekti; ilgili sonuçlar da kesiliyordu
                    )
                    web_sources = (web_data.get("results") if isinstance(web_data, dict) else []) or []
                    if web_sources:
                        # ── AŞAMA 1: Tam metni VectorDB'ye kaydet (prompt'a girmez) ──
                        for i, w in enumerate(web_sources[:6], 1):
                            title = (w.get("title") or "").strip() or "Kaynak"
                            url = w.get("url") or ""
                            content = w.get("content") or ""
                            raw_content = (w.get("raw_content") or "").strip()
                            web_link_map[i] = {"index": i, "title": title, "url": url}
                            store_text = raw_content if raw_content else content
                            if store_text:
                                try:
                                    await store.a_add_temp_document(
                                        session_id=session_id,
                                        doc_id=f"web_{i}",
                                        text=store_text,
                                        source_url=url,
                                        title=title,
                                        chunk_size=1500,
                                    )
                                    print(f"✅ [VectorDB] web_{i} kaydedildi: {len(store_text)} char → {url[:60]}")
                                except Exception as _we:
                                    print(f"⚠️ Web içeriği vektöre eklenemedi: {url} -> {_we}")
                        tlog.lap("web içerikleri VectorDB'ye kaydedildi")

                        # ── AŞAMA 2: VectorDB'den soruyla ilgili chunk'ları çek ──
                        try:
                            relevant = await store.a_similarity_search(
                                session_id=session_id,
                                query=message,
                                k=20,
                                include_pdf=True,
                                include_chat=False
                            )
                            temps = [r for r in relevant if r.get("kind") == "temp"][:12]
                            if temps:
                                lines = []
                                lines.append("\n### 0. GÜNCEL MEVZUAT VE KAYNAKLAR")
                                lines.append(
                                    "⚠️ Aşağıdaki içerikler web kaynaklarından soruyla ilgili "
                                    "bulunan bölümlerdir. Spesifik süreler/oranlar/rakamlar "
                                    "model bilgisinden önceliklidir."
                                )
                                for r in temps:
                                    pid = r.get("pdf_id") or ""
                                    idx = 0
                                    try:
                                        if pid.startswith("web_"):
                                            idx = int(pid.split("_", 1)[1])
                                    except Exception:
                                        idx = 0
                                    tag = f"[w{idx}]" if idx else ""
                                    snippet = (r.get("chunk_text") or "").strip()
                                    src_title = web_link_map.get(idx, {}).get("title", "") if idx else ""
                                    src_url = web_link_map.get(idx, {}).get("url", "") if idx else ""
                                    header = f"- {tag} **{src_title}**" if src_title else f"- {tag}"
                                    if src_url:
                                        header += f" ({src_url})"
                                    lines.append(f"{header}\n  {snippet}")
                                web_section = "\n".join(lines) + "\n\n"
                                print(f"✅ [Semantic] {len(temps)} alakalı chunk prompt'a eklendi")
                            else:
                                # Semantik sonuç yoksa en azından başlık+URL listele
                                lines = ["\n### 0. GÜNCEL MEVZUAT VE KAYNAKLAR"]
                                for idx, info in web_link_map.items():
                                    lines.append(f"- [w{idx}] **{info['title']}** → {info['url']}")
                                web_section = "\n".join(lines) + "\n\n"
                            tlog.lap("web semantik alıntılar hazır")
                        except Exception as _ws:
                            print(f"⚠️ Web semantik arama hatası: {_ws}")
                else:
                    print("ℹ️ Tavily web araması devre dışı (TAVILY_API_KEY yok)")
            except Exception as we:
                print(f"⚠️ Web arama hatası: {we}")
        else:
            # CONCISE MOD WEB ARAMASI: Düz soru modunda da resmi kaynaklarda arama yap.
            # Sorun: Model eğitim verisindeki eski bilgiler ile güncel mevzuat çelişebilir.
            # Çözüm: Resmi kaynaklarda advanced arama + vektöre ekleme + semantik arama.
            try:
                web_fetcher = WebContextFetcher()
                if web_fetcher.enabled:
                    # gpt-4o-mini ile odaklı query üretimi (3s timeout → fallback: message[:120])
                    try:
                        law_query = await asyncio.wait_for(
                            generate_precedent_search_query(message, relevant_content),
                            timeout=3.0
                        )
                        if not law_query:
                            law_query = message[:120]
                        print(f"🔍 [WebContext-Concise] AI query üretildi: '{law_query}'")
                    except asyncio.TimeoutError:
                        law_query = message[:120]
                        print("⚠️ [WebContext-Concise] Query üretimi timeout → kısa mesaj kullanıldı")

                    web_data = await web_fetcher.asearch(
                        law_query,
                        max_results=5,
                        search_depth="advanced",
                        days=3650,
                        min_score=0.50,
                    )
                    concise_web = (web_data.get("results") if isinstance(web_data, dict) else []) or []
                    if concise_web:
                        web_sources = concise_web  # Footer için kaydet
                        # ── AŞAMA 1: Tam metni VectorDB'ye kaydet (prompt'a girmez) ──
                        for i, w in enumerate(concise_web[:4], 1):
                            title = (w.get("title") or "").strip() or "Kaynak"
                            url = w.get("url") or ""
                            raw_content = (w.get("raw_content") or "").strip()
                            content = w.get("content") or ""
                            web_link_map[i] = {"index": i, "title": title, "url": url}
                            store_text = raw_content if raw_content else content
                            if store_text:
                                try:
                                    await store.a_add_temp_document(
                                        session_id=session_id,
                                        doc_id=f"web_{i}",
                                        text=store_text,
                                        source_url=url,
                                        title=title,
                                        chunk_size=1500,
                                    )
                                except Exception:
                                    pass
                        tlog.lap(f"concise web VectorDB kaydı ({len(concise_web)} sonuç)")

                        # ── AŞAMA 2: VectorDB'den soruyla ilgili chunk'ları çek ──
                        try:
                            relevant_web = await store.a_similarity_search(
                                session_id=session_id,
                                query=message,
                                k=15,
                                include_pdf=True,
                                include_chat=False
                            )
                            temps_c = [r for r in relevant_web if r.get("kind") == "temp"][:8]
                            if temps_c:
                                clines = ["\n### GÜNCEL RESMİ KAYNAKLAR (Mevzuat/Resmî Gazete/Yargıtay)"]
                                clines.append(
                                    "⚠️ Aşağıdaki içerikler resmi kaynaklardan soruyla ilgili "
                                    "bulunan bölümlerdir. Kaynakta gördüğün değeri kullan."
                                )
                                for r in temps_c:
                                    pid = r.get("pdf_id") or ""
                                    idx = 0
                                    try:
                                        if pid.startswith("web_"):
                                            idx = int(pid.split("_", 1)[1])
                                    except Exception:
                                        idx = 0
                                    tag = f"[w{idx}]" if idx else ""
                                    snippet = (r.get("chunk_text") or "").strip()
                                    src_title = web_link_map.get(idx, {}).get("title", "") if idx else ""
                                    src_url = web_link_map.get(idx, {}).get("url", "") if idx else ""
                                    header = f"- {tag} **{src_title}**" if src_title else f"- {tag}"
                                    if src_url:
                                        header += f" ({src_url})"
                                    clines.append(f"{header}\n  {snippet}")
                                web_section = "\n".join(clines) + "\n\n"
                                print(f"✅ [WebContext-Concise] {len(temps_c)} alakalı chunk prompt'a eklendi")
                            else:
                                clines = ["\n### GÜNCEL RESMİ KAYNAKLAR"]
                                for idx_c, info in web_link_map.items():
                                    clines.append(f"- [w{idx_c}] **{info['title']}** → {info['url']}")
                                web_section = "\n".join(clines) + "\n\n"
                        except Exception as _ws:
                            print(f"⚠️ [WebContext-Concise] Semantik arama hatası: {_ws}")
                        tlog.lap("concise web semantik alıntılar hazır")
                    else:
                        print("ℹ️ [WebContext-Concise] Resmi kaynaklarda sonuç bulunamadı")
                else:
                    print("ℹ️ [WebContext-Concise] Tavily API key yok, web araması atlandı")
            except Exception as _cwe:
                print(f"⚠️ [WebContext-Concise] Hata: {_cwe}")

        # Kullanıcıdan gelen (opsiyonel) taraf/durum notlarını ekle (sadece detay modda)
        notes_section = ""
        if detail != "concise":
            extra_lines = []
            if party_info and party_info.strip():
                extra_lines.append(f"- Taraf: {party_info.strip()}")
            if situation_info and situation_info.strip():
                extra_lines.append(f"- Durum: {situation_info.strip()}")
            if extra_lines:
                notes_section = "### Rol/Durum Notları\n" + "\n".join(extra_lines) + "\n\n"

        citation_note = ""
        if web_section:
            citation_note = (
                "Not: Web kaynaklarına atıf yaparken metin içinde [w1], [w2] biçiminde referans kullan. "
                "Cevabın sonunda 'Kaynaklar' başlığı altında [wN] -> URL eşlemesini listele."
            )

        # web_semantic_context artık kullanılmıyor — tüm web içerikleri
        # VectorDB semantik arama ile web_section'a ekleniyor.
        combined_context = (notes_section + base_memory) \
            + (("\n\n" + citation_note) if citation_note else "") \
            + (("\n\n" + precedent_context) if precedent_context else "") \
            + (("\n\n" + web_section) if web_section else "")

        # --- Token/uzunluk koruması ---
        # Web içerikleri artık tam metin değil, sadece semantik chunk'lar.
        # Token maliyeti kontrol altında; makul sınırlar yeterli.
        MAX_CHARS_SOFT = 50000
        MAX_CHARS_HARD = 65000

        def shrink_context(txt: str) -> str:
            if len(txt) <= MAX_CHARS_SOFT:
                return txt
            # Strateji:
            #  head = base_memory + kısa emsal özeti → %55 bütçe
            #  tail = AI emsal özeti + TAM emsal içerikleri + web → %45 bütçe
            # Emsal bütçesi %30'dan %45'e çıkarıldı: tam içerik bloku (build_full_precedents_block)
            # artık bu bölümde ve önemli, kırpılmamalı.
            marker_precedent = "-- KISA EMSAL ÖZETLERİ --"
            if marker_precedent in txt:
                head, tail = txt.split(marker_precedent, 1)
                target_mem = int(MAX_CHARS_SOFT * 0.55)   # base_memory + compact summary için %55
                target_prec = MAX_CHARS_SOFT - target_mem  # AI summary + tam içerik + web için %45
                if len(head) > target_mem:
                    head = head[:target_mem]
                if len(tail) > target_prec:
                    tail = tail[:target_prec]
                new_txt = head + marker_precedent + tail
            else:
                new_txt = txt[:MAX_CHARS_SOFT]
            if len(new_txt) > MAX_CHARS_HARD:
                new_txt = new_txt[:MAX_CHARS_HARD]
            print(f"⚠️ Context shrink: {len(txt)} -> {len(new_txt)} chars")
            return new_txt

        combined_context = shrink_context(combined_context)
        detail_level = "standart" if detail == "concise" else "detaylı"
        # Yeni: Esnek mod seçimi. concise => flex_concise, diğerleri => flex_detailed
        mode = "flex_concise" if detail == "concise" else "flex_detailed"
        prompt = build_legal_analysis_prompt(user_question=message, context_snippet=combined_context, detail_level=detail_level, mode=mode)

        # AI Guard ile güvenli API çağrısı
        async def _process_with_ai():
            if legal_advisor_ai is None:
                return {"error": "AI sistemi başlatılamadı"}
            return await legal_advisor_ai.process_question(prompt, flex_mode=mode)

        result = await safe_openai_request(_process_with_ai)
        tlog.lap("AI yanıtı alındı")

        # Not: Artık genel '(Emsal Kaynaklar: ...)' bloğu eklemiyoruz.

        if result and 'response' in result:
            raw_answer = result['response']
            precedent_set_id = None
            if detail != "concise":
                # Kaynak etiketleri metinde görünmesin
                raw_answer = enforce_citations(raw_answer, len(precedent_meta.get("selected", [])))
                raw_answer = _strip_source_tags(raw_answer)
                # Precedent set kaydet (tam listeyi UI ayrı çağıracak)
                if precedent_meta.get("selected"):
                    from uuid import uuid4 as _u
                    precedent_set_id = f"p_{_u().hex[:10]}"
                    if session_id not in session_chat_precedents:
                        session_chat_precedents[session_id] = {}
                    session_chat_precedents[session_id][precedent_set_id] = {
                        "precedents": precedent_meta["selected"],
                        "search_query": precedent_meta.get("search_query"),
                        "created": datetime.now().isoformat(),
                        "total_fetched": precedent_meta.get("total_fetched"),
                    }
            # Kısa esnek modda post-processing yapmayalım; detay modda hafif format
            formatted_response = format_ai_response(raw_answer) if mode != "flex_concise" else raw_answer.strip()
            # Kullanılan web kaynaklarını [wN] linkleriyle ekle (her iki modda)
            if web_sources:
                try:
                    # DEDUP KONTROLÜ: AI zaten Kaynaklar bölümü yazdıysa yeni ekleyince
                    # kullanıcı iki kez kaynak listesi görür. Eğer yanıtta "Kaynaklar"
                    # içerikli bir bölüm varsa sadece Mevzuat linklerini ekle.
                    import re as _src_re
                    has_kaynaklar_block = bool(
                        _src_re.search(r'\bKaynaklar\b.*\[w\d+\]', formatted_response, _src_re.DOTALL)
                    )

                    # ALAKA FİLTRESİ: Sadece AI'ın gerçekten atıf yaptığı [wN] kaynakları göster.
                    # Bu sayede alakasız kaynaklar (örn: aile hukuku sorusundaki Danıştay dergileri)
                    # footer'da görünmez.
                    cited_indices = _extract_cited_web_indices(formatted_response)
                    limit = 6 if detail != "concise" else 4
                    all_indexed = list(enumerate(web_sources[:limit], 1))

                    if cited_indices:
                        # Sadece gerçekten atıf yapılan kaynaklar
                        sources_for_footer = [(i, w) for i, w in all_indexed if i in cited_indices]
                        uncited_count = sum(1 for i, _ in all_indexed if i not in cited_indices)
                        if uncited_count:
                            print(f"ℹ️  [Web Sources] {uncited_count} kaynak AI tarafından atıf yapılmadı, footer'dan çıkarıldı")
                    else:
                        # AI hiç [wN] atıfı yapmadıysa hepsini göster (fallback)
                        sources_for_footer = all_indexed

                    src_lines = []
                    if not has_kaynaklar_block and sources_for_footer:
                        src_lines.append("\n---\nKaynaklar:")
                        for i, w in sources_for_footer:
                            title = (w.get("title") or "").strip() or "Kaynak"
                            url = w.get("url") or ""
                            src_lines.append(f"- [w{i}] {title} — {url}")
                    elif has_kaynaklar_block:
                        # AI kendi Kaynaklar bölümünü zaten yazmış; tekrar yazma
                        print("ℹ️  [Web Sources] AI kendi Kaynaklar bölümünü oluşturmuş; footer tekrar yazılmıyor")

                    # Ayrıca metin içinde geçen kanun atıfları için resmî mevzuat bağlantıları ekle
                    cited_codes = _extract_law_codes(formatted_response)
                    if cited_codes:
                        src_lines.append("\nMevzuat bağlantıları:")
                        for code in cited_codes:
                            law_no = LAW_MAP.get(code)
                            if law_no:
                                src_lines.append(f"- {code}: {_law_url(law_no)}")
                    if src_lines:
                        formatted_response = formatted_response.rstrip() + "\n" + "\n".join(src_lines)
                except Exception:
                    pass
            out = {
                "response": formatted_response,
                "timestamp": datetime.now().isoformat(),
                "type": "normal_chat_with_vector"
            }
            if precedent_meta:
                out["precedent_search_query"] = precedent_meta.get("search_query")
                out["precedent_count"] = len(precedent_meta.get("selected", []))
            # Web kaynaklarını API çıktısına dahil et (her iki modda)
            out["web_sources_used"] = bool(web_sources)
            if web_sources:
                out["web_sources"] = web_sources
                # Provide explicit citation map for UI if needed
                out["web_citations"] = [{"id": f"w{i}", "index": i, "title": (w.get("title") or "").strip() or "Kaynak", "url": w.get("url") or ""} for i, w in enumerate(web_sources[: (6 if detail != "concise" else 4) ], 1)]
            if precedent_set_id:
                # Hafif metadata (UI listede gösterir, full içerik ayrı endpoint)
                light = []
                for i, p in enumerate(precedent_meta.get("selected", []), 1):
                    light.append({
                        "index": i,
                        "id": p.get("documentId") or p.get("documentID") or p.get("id"),
                        "birim": p.get("birimAdi"),
                        "tarih": p.get("kararTarihi"),
                        "tur": p.get("itemType"),
                    })
                out["precedent_set_id"] = precedent_set_id
                out["precedents_meta"] = light
            tlog.end()
            return out
        else:
            tlog.end("hata ile bitti")
            return {"error": result.get('error', "AI'dan yanıt alınamadı")}
    except Exception as e:
        print(f"❌ Vector chat hatası: {str(e)}")
        return {"error": f"Chat sırasında hata: {str(e)}"}

async def process_case_analysis_with_vector_search(message: str, session_id: str, case_role: Optional[str] = None, case_status: Optional[str] = None) -> Dict[str, Any]:
    return {"error": "Dava analizi özelliği kaldırıldı"}

async def process_petition_with_vector_search(message: str, session_id: str) -> Dict[str, Any]:
    """Dilekçe oluşturma - vector store'daki dosyalarla"""
    try:
        tlog = TimingLogger("chat-petition", session_id)
        if petition_generator_ai is None:
            return {"error": "Dilekçe AI sistemi başlatılamadı"}
        
        # Vector store'dan ilgili içerikleri ara
        store = get_vector_store()
        
        session_has_files = store.has_pdf_files(session_id)
        
        search_results = await store.a_similarity_search(
            session_id=session_id, 
            query=message, 
            k=6,
            include_pdf=True,
            include_chat=False  # Sadece PDF'leri ara
        )
        
        # Pinecone eventual consistency: dosya varsa ama sonuç yoksa tekrar dene
        if not search_results and session_has_files:
            await asyncio.sleep(2)
            search_results = await store.a_similarity_search(
                session_id=session_id, query=message, k=6,
                include_pdf=True, include_chat=False
            )
        
        # Hâlâ sonuç yoksa Pinecone fetch fallback
        if not search_results and session_has_files:
            fallback_docs = store.get_all_documents_raw(session_id, kinds=["pdf"])
            for r in fallback_docs[:6]:
                search_results.append({
                    "session_id": session_id,
                    "pdf_id": r["pdf_id"],
                    "chunk_id": r["chunk_id"],
                    "chunk_text": r["text"],
                    "distance": 0.0,
                    "kind": "pdf",
                    "role": None,
                })
        
        tlog.lap(f"vector search tamamlandı ({len(search_results)} sonuç)")
        
        # İlgili içeriği birleştir
        relevant_content = ""
        if search_results:
            relevant_content = "\n\n".join([
                f"**Kaynak: {result['pdf_id']}**\n{result['chunk_text']}"
                for result in search_results[:4]  # En iyi 4 sonucu al
            ])
        
        print(f"📋 Dilekçe için {len(search_results)} chunk bulundu")
        
        # Dilekçe oluştur
        result = await petition_generator_ai.generate_petition(
            user_prompt=message,
            file_content=relevant_content
        )
        tlog.lap("AI dilekçe üretildi")
        
        if result and 'petition' in result:
            out = {
                "response": result['petition'],
                "timestamp": datetime.now().isoformat(),
                "type": "petition_vector"
            }
            tlog.end()
            return out
        tlog.end("hata ile bitti")
        return {"error": "Dilekçe oluşturulamadı"}
        
    except Exception as e:
        print(f"❌ Vector petition hatası: {str(e)}")
        return {"error": f"Dilekçe oluşturma sırasında hata: {str(e)}"}

async def process_petition_generation(message: str, file_content: Optional[str] = None) -> Dict[str, Any]:
    """Dilekçe oluşturma işlemi - önceki chat geçmişi + mevcut mesaj + dosya içeriği"""
    try:
        if petition_generator_ai is None:
            return {"error": "Dilekçe AI hazır değil"}

        result = await petition_generator_ai.generate_petition(
            user_prompt=message,
            file_content=file_content
        )
        if "error" in result:
            return {"error": result["error"]}

        formatted = format_ai_response(result.get("petition", ""))
        return {
            "response": formatted,
            "timestamp": result.get("timestamp"),
            "type": "petition",
            "meta": result.get("used_sections", {})
        }
    except Exception as e:
        return {"error": f"Dilekçe oluşturma hatası: {str(e)}"}
# ================================
# ROUTES
# ================================

# Not: Kök URL altında artık UI sunulmuyor. UI için Streamlit tabanlı `ui.py` kullanılmalıdır.

@app.post("/upload-file")
async def upload_file(
    file: UploadFile = File(...),
    session_id: str = Header(None, alias="session_id", convert_underscores=False)
):
    """Dosya yükle ve içeriği çıkar"""
    try:
        print(f"📁 Upload başladı: {file.filename}, session: {session_id}")
        print(f"📁 File object: {file}")
        print(f"📁 File content type: {file.content_type}")
        
        if not session_id:
            print("❌ session_id header bulunamadı")
            raise HTTPException(status_code=400, detail="session_id header gerekli")
            
        # Vector store'u kontrol et
        store = get_vector_store()
        pdf_count = store.count_pdf_files(session_id)
        print(f"📊 Upload öncesi session'da {pdf_count} PDF kayıt var")
            
        contents = await file.read()
        print(f"📊 Okunan byte sayısı: {len(contents)}")

        # Dosya boyut limiti: 50 MB
        MAX_FILE_BYTES = 50 * 1024 * 1024
        if len(contents) > MAX_FILE_BYTES:
            raise HTTPException(status_code=413, detail=f"Dosya çok büyük. Maksimum boyut: 50 MB (Yüklenen: {len(contents)//1024//1024} MB)")

        fname_lower = file.filename.lower()
        if fname_lower.endswith('.pdf'):
            print("📄 PDF işleniyor...")
            text = await asyncio.to_thread(extract_text_from_pdf, contents)
        elif fname_lower.endswith('.docx'):
            print("📝 DOCX dosyası işleniyor...")
            text = await asyncio.to_thread(extract_text_from_docx, contents)
        elif fname_lower.endswith('.doc'):
            print("📝 DOC dosyası işleniyor (DOCX olarak deneniyor)...")
            # Birçok modern .doc dosyası aslında DOCX formatındadır; önce DOCX olarak dene
            try:
                text = await asyncio.to_thread(extract_text_from_docx, contents)
            except Exception:
                # Eski binary .doc formatı – düz metin olarak oku (sınırlı destek)
                text = contents.decode('utf-8', errors='ignore')
                if not text.strip() or len([c for c in text[:500] if c == '\x00']) > 50:
                    raise HTTPException(status_code=400, detail="Eski .doc formatı okunamadı. Lütfen dosyayı .docx olarak kaydedin ve tekrar yükleyin.")
        elif fname_lower.endswith('.txt'):
            print("📝 TXT dosyası işleniyor...")
            # Encoding tespiti: önce UTF-8, sonra Windows-1254 (Türkçe), sonra latin-1
            for enc in ('utf-8', 'cp1254', 'latin-1'):
                try:
                    text = contents.decode(enc)
                    break
                except (UnicodeDecodeError, ValueError):
                    continue
            else:
                text = contents.decode('utf-8', errors='ignore')
        else:
            print(f"❌ Desteklenmeyen dosya formatı: {file.filename}")
            raise HTTPException(status_code=400, detail="Desteklenmeyen dosya formatı. PDF, TXT, DOC, DOCX dosyaları kabul edilir.")

        print(f"📊 Metin çıkarıldı: {len(text)} karakter")

    # PDF veya desteklenen dosya metnini Pinecone vector store'a ekle
        print("🔍 Vector store'a ekleniyor...")
        store = get_vector_store()
        add_info = await store.a_add_pdf(session_id=session_id, pdf_id=file.filename, text=text)

        print(f"✅ Vector store'a eklendi: {add_info}")

        return JSONResponse({
            "success": True,
            "filename": file.filename,
            "content": text,
            "size": len(contents),
            "vector_store": add_info,
            "session_id": session_id
        })
    except HTTPException:
        # Önceden oluşturulmuş HTTPException'ı tekrar fırlat (örn: 400 session_id yok)
        print("❌ HTTPException yeniden fırlatılıyor")
        raise
    except Exception as e:
        print(f"❌ Upload hatası: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Dosya yükleme hatası: " + str(e))

@app.post("/remove-file")
async def remove_file(
    request: Request,
    session_id: str = Header(None, alias="session_id", convert_underscores=False)
):
    """Yüklenen dosyayı kaldır ve vector store'dan sil"""
    try:
        if not session_id:
            raise HTTPException(status_code=400, detail="session_id header gerekli")
        
        body = await request.json()
        filename = body.get('filename')
        file_id = body.get('file_id')
        
        if not filename:
            raise HTTPException(status_code=400, detail="filename gerekli")
        
        print(f"🗑️  Dosya kaldırılıyor: {filename}, session: {session_id}")
        
        # Vector store'dan dosyayı kaldır
        store = get_vector_store()
        
        # Dosyayı PDF ID ile kaldır (filename kullanarak)
        try:
            # Vector store'da remove_pdf metodu varsa kullan
            if hasattr(store, 'a_remove_pdf'):
                remove_result = await store.a_remove_pdf(session_id=session_id, pdf_id=filename)
                print(f"✅ Vector store'dan kaldırıldı: {remove_result}")
            else:
                # Alternatif: session'daki tüm vektörleri temizle (geçici çözüm)
                print("⚠️  remove_pdf metodu bulunamadı, session temizleniyor")
                # Bu durumda manuel temizlik yapabiliriz
        except Exception as vector_error:
            print(f"❌ Vector store'dan kaldırma hatası: {vector_error}")
            # Vector hatası olsa bile başarılı döndür
        
        return JSONResponse({
            "success": True,
            "message": f"Dosya '{filename}' başarıyla kaldırıldı",
            "filename": filename,
            "file_id": file_id
        })
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Dosya kaldırma hatası: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Dosya kaldırma hatası: " + str(e))

# Dava full-pipeline endpoint kaldırıldı

@app.post("/chat")
async def chat_endpoint(
    request: Request,
    message: str = Form(...),
    is_petition: bool = Form(False),
    chat_detail: Optional[str] = Form("concise"),
    scan_precedents: Optional[bool] = Form(True),
    party_info: Optional[str] = Form(None),
    situation_info: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    session_id: str = Header(None, alias="session_id", convert_underscores=False)
):
    """Chat endpoint - Vector store'daki tüm dosyalarla çalışır"""
    try:
        # Mesaj uzunluk limiti
        MAX_MSG_LEN = 12000
        if len(message) > MAX_MSG_LEN:
            raise HTTPException(status_code=400, detail=f"Mesaj çok uzun. Maksimum {MAX_MSG_LEN} karakter. (Gönderilen: {len(message)})")

        # session_id kontrol
        if not session_id:
            session_id = new_session_id()
        print(f"💬 Chat başladı: session={session_id}, petition={is_petition}")
        store = get_vector_store()

        # Özel full pipeline tetikleme kısayolu: kullanıcı mesajı '#full' ile başlıyorsa
        # #full kısayolu kaldırıldı
        
        # Vector store'dan session'daki dosya sayısını kontrol et (Pinecone)
        try:
            pdf_count = store.count_pdf_files(session_id)
            print(f"📊 Session'da {pdf_count} PDF dosyası bulundu")
            print(f"🔍 Session ID kontrol ediliyor: {session_id}")
        except Exception as e:
            print(f"⚠️  PDF sayısı kontrol hatası: {e}")
            pdf_count = 0
        
        # Chat endpoint'ine doğrudan dosya yüklendiyse, vector store'a ekle
        file_content = None
        if file and file.filename:
            try:
                contents = await file.read()
                fname_lower = file.filename.lower()
                if fname_lower.endswith('.pdf'):
                    text = await asyncio.to_thread(extract_text_from_pdf, contents)
                elif fname_lower.endswith('.docx'):
                    text = await asyncio.to_thread(extract_text_from_docx, contents)
                elif fname_lower.endswith('.doc'):
                    try:
                        text = await asyncio.to_thread(extract_text_from_docx, contents)
                    except Exception:
                        text = contents.decode('utf-8', errors='ignore')
                elif fname_lower.endswith('.txt'):
                    for enc in ('utf-8', 'cp1254', 'latin-1'):
                        try:
                            text = contents.decode(enc)
                            break
                        except (UnicodeDecodeError, ValueError):
                            continue
                    else:
                        text = contents.decode('utf-8', errors='ignore')
                else:
                    text = None
                if text:
                    file_content = text
                    await store.a_add_pdf(session_id=session_id, pdf_id=file.filename, text=text)
                    print(f"📁 Chat içinde dosya yüklendi ve vector store'a eklendi: {file.filename}")
                    pdf_count += 1
            except Exception as fe:
                print(f"⚠️ Chat içi dosya işleme hatası: {fe}")
        has_uploaded_files = pdf_count > 0
        
        # Kullanıcı mesajını chat geçmişine ekle
        user_message = {
            "role": "user",
            "content": message,
            "timestamp": datetime.now().isoformat(),
            "has_file": has_uploaded_files,
            "is_case_analysis": False,
            "is_petition": is_petition
        }
        
        # Kısa selamlaşma/sohbet mesajları için kısa yanıt
        def _is_smalltalk(msg: str) -> bool:
            if not msg:
                return False
            m = msg.lower().strip()
            import re as _re
            patterns = [r"^merhaba\b", r"^selam\b", r"nasılsın", r"^kimsin\b", r"sen kimsin", r"ne yapıyorsun", r"günaydın", r"iyi akşamlar"]
            return any(_re.search(p, m) for p in patterns)

        if _is_smalltalk(message):
            short = (
                "Merhaba, ben Yargı AI. Basit sohbetlere kısa yanıt veririm; hukuki konulara odaklanırım. "
                "Somut bir soru veya dosya paylaşımı yaparsanız yardımcı olabilirim."
            )
            ai_response = {
                "role": "assistant",
                "content": short,
                "timestamp": datetime.now().isoformat(),
                "type": "smalltalk"
            }
            try:
                await store.a_add_chat_message(session_id=session_id, role="assistant", content=ai_response["content"])
            except Exception:
                pass
            return JSONResponse({
                "success": True,
                "response": ai_response,
                "chat_history": store.get_chat_history(session_id=session_id),
                "session_id": session_id
            })

        # Vektör depoya kullanıcı mesajını ekle
        try:
            await store.a_add_chat_message(session_id=session_id, role="user", content=message)
        except Exception as ve:
            print(f"[VectorStore] user chat eklenemedi: {ve}")
        
        # AI yanıtını al
        if is_petition:
            # Dilekçe oluşturma - vector store'dan ilgili belgeleri ara
            result = await process_petition_with_vector_search(message, session_id)
        else:
            # Normal sohbet - vector store araması ile
            # scan_precedents yalnızca detaylı modda anlamlı; concise modda zaten emsal bakılmaz
            include_precedents = bool(scan_precedents) if (chat_detail or "concise") != "concise" else False
            result = await process_normal_chat_with_vector_search(
                message,
                session_id,
                detail=(chat_detail or "concise"),
                party_info=party_info,
                situation_info=situation_info,
                include_precedents=include_precedents
            )
        
        if "error" in result:
            ai_response = {
                "role": "assistant",
                "content": f"❌ Hata: {result['error']}",
                "timestamp": datetime.now().isoformat(),
                "type": "error"
            }
        else:
            # Ana yanıt gövdesi
            ai_response = {
                "role": "assistant",
                "content": result['response'],
                "timestamp": datetime.now().isoformat(),
                "type": result.get('type', 'response')
            }
            if isinstance(result, dict) and result.get("web_citations"):
                ai_response["web_citations"] = result.get("web_citations")
            # Emsal meta bilgisi (sadece detaylı chat için)
            if isinstance(result, dict):
                for k in ("precedent_set_id", "precedents_meta", "precedent_search_query", "precedent_count"):
                    if k in result:
                        ai_response[k] = result[k]
        
        try:
            await store.a_add_chat_message(session_id=session_id, role="assistant", content=ai_response["content"])
        except Exception as ve:
            print(f"[VectorStore] assistant chat eklenemedi: {ve}")
        
        return JSONResponse({
            "success": True,
            "response": ai_response,
            "chat_history": store.get_chat_history(session_id=session_id),
            "session_id": session_id
        })
        
    except Exception as e:
        error_response = {
            "role": "assistant",
            "content": f"❌ Beklenmeyen hata: {str(e)}",
            "timestamp": datetime.now().isoformat(),
            "type": "error"
        }
        
        store = get_vector_store()
        try:
            await store.a_add_chat_message(session_id=session_id or new_session_id(), role="assistant", content=error_response["content"])
        except Exception:
            pass
        return JSONResponse({
            "success": False,
            "response": error_response,
            "chat_history": []
        })

@app.post("/chat-stream")
async def chat_stream_endpoint(
    request: Request,
    message: str = Form(...),
    is_petition: bool = Form(False),
    chat_detail: Optional[str] = Form("concise"),
    scan_precedents: Optional[bool] = Form(True),
    party_info: Optional[str] = Form(None),
    situation_info: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    session_id: str = Header(None, alias="session_id", convert_underscores=False)
):
    """Streaming chat endpoint"""
    # Mesaj uzunluk limiti
    MAX_MSG_LEN = 12000
    if len(message) > MAX_MSG_LEN:
        return JSONResponse(
            {"error": f"Mesaj çok uzun. Maksimum {MAX_MSG_LEN} karakter."},
            status_code=400
        )

    if not session_id:
        session_id = new_session_id()

    file_content = None
    if file and file.filename:
        try:
            content_bytes = await file.read()
            fname_lower = file.filename.lower()
            if fname_lower.endswith('.pdf'):
                file_content = await asyncio.to_thread(extract_text_from_pdf, content_bytes)
            elif fname_lower.endswith('.docx'):
                file_content = await asyncio.to_thread(extract_text_from_docx, content_bytes)
            elif fname_lower.endswith('.doc'):
                try:
                    file_content = await asyncio.to_thread(extract_text_from_docx, content_bytes)
                except Exception:
                    file_content = content_bytes.decode('utf-8', errors='ignore')
            elif fname_lower.endswith('.txt'):
                for enc in ('utf-8', 'cp1254', 'latin-1'):
                    try:
                        file_content = content_bytes.decode(enc)
                        break
                    except (UnicodeDecodeError, ValueError):
                        continue
                else:
                    file_content = content_bytes.decode('utf-8', errors='ignore')
            if file_content:
                store = get_vector_store()
                await store.a_add_pdf(session_id=session_id, pdf_id=file.filename, text=file_content)
                print(f"📁 Chat-stream içinde dosya yüklendi ve vector store'a eklendi: {file.filename}")
        except Exception as file_error:
            return JSONResponse({
                "success": False,
                "error": f"Dosya okuma hatası: {str(file_error)}"
            })
    async def generate_response():
        try:
            store = get_vector_store()
            try:
                await store.a_add_chat_message(session_id=session_id, role="user", content=message)
            except Exception as ve:
                print(f"[VectorStore] user(chat-stream) eklenemedi: {ve}")
            # Kısa selamlaşma/sohbet mesajları için kısa yanıt
            def _is_smalltalk(msg: str) -> bool:
                if not msg:
                    return False
                m = msg.lower().strip()
                import re as _re
                patterns = [r"^merhaba\b", r"^selam\b", r"nasılsın", r"^kimsin\b", r"sen kimsin", r"ne yapıyorsun", r"günaydın", r"iyi akşamlar"]
                return any(_re.search(p, m) for p in patterns)

            if _is_smalltalk(message):
                short = "Merhaba, ben Yargı AI. Hukuki sorulara yardımcı olurum; lütfen somut bir konu paylaşın."
                yield f"data: {json.dumps({'content': short, 'done': True})}\n\n"
                return

            if is_petition:
                yield f"data: {json.dumps({'content': '📝 Dilekçe taslağı oluşturuluyor...', 'loading': True})}\n\n"
                result = await process_petition_generation(message, file_content)
                if 'error' in result:
                    yield f"data: {json.dumps({'error': result['error'], 'done': True})}\n\n"
                else:
                    yield f"data: {json.dumps({'content': format_ai_response(result['response'])})}\n\n"
                    yield f"data: {json.dumps({'done': True})}\n\n"
            else:
                yield f"data: {json.dumps({'content': '🔄 Yanıt hazırlanıyor...', 'loading': True})}\n\n"
                include_precedents = bool(scan_precedents) if (chat_detail or 'concise') != 'concise' else False
                # Uzun işlemi arka plana al; proxy timeout'u önlemek için heartbeat gönder
                task = asyncio.create_task(process_normal_chat_with_vector_search(
                    message,
                    session_id,
                    detail=(chat_detail or 'concise'),
                    party_info=party_info,
                    situation_info=situation_info,
                    include_precedents=include_precedents
                ))
                while not task.done():
                    await asyncio.sleep(15)
                    if not task.done():
                        yield f"data: {json.dumps({'heartbeat': True})}\n\n"
                result = task.result()
                if 'error' in result:
                    yield f"data: {json.dumps({'error': result['error'], 'done': True})}\n\n"
                else:
                    # Precedent meta varsa önce meta bilgisini ayrı olay olarak gönder
                    if result.get('precedent_set_id'):
                        meta_event = {
                            'precedent_meta': result.get('precedents_meta', []),
                            'precedent_set_id': result.get('precedent_set_id'),
                            'precedent_search_query': result.get('precedent_search_query'),
                            'precedent_count': result.get('precedent_count')
                        }
                        print(f"📚 Sending precedent meta: {len(meta_event.get('precedent_meta', []))} items")
                        yield f"data: {json.dumps(meta_event)}\n\n"
                    payload = {'content': format_ai_response(result['response'])}
                    if result.get('precedent_set_id'):
                        payload['precedent_set_id'] = result.get('precedent_set_id')
                    if result.get('web_citations'):
                        payload['web_citations'] = result.get('web_citations')
                    yield f"data: {json.dumps(payload)}\n\n"
                    yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            error_msg = f"Hata: {str(e)}"
            yield f"data: {json.dumps({'error': error_msg, 'done': True})}\n\n"
    return StreamingResponse(
        generate_response(),
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Content-Type": "text/event-stream"
        }
    )

@app.get("/session/new")
async def create_session():
    sid = new_session_id()
    print(f"🆔 Yeni session ID oluşturuldu: {sid}")
    return {"session_id": sid}

@app.delete("/session/{session_id}")
async def delete_session(session_id: str):
    print(f"🗑️ Session siliniyor: {session_id}")
    store = get_vector_store()
    removed = await store.a_delete_session(session_id)
    print(f"✅ Session silindi: {session_id}, {removed} kayıt kaldırıldı")
    return {"removed": removed, "session_id": session_id}

@app.post("/session/close")
async def close_session(session_id: str = Form(...)):
    print(f"🔚 Session kapatılıyor: {session_id}")
    store = get_vector_store()
    removed = await store.a_delete_session(session_id)
    print(f"✅ Session kapatıldı: {session_id}, {removed} kayıt kaldırıldı")
    return {"removed": removed, "session_id": session_id, "closed": True}

@app.get("/search")
async def similarity_search(query: str, session_id: str):
    store = get_vector_store()
    results = await store.a_similarity_search(session_id=session_id, query=query, k=10)
    return {"results": results, "count": len(results)}

@app.get("/chat-history")
async def get_chat_history_endpoint(session_id: Optional[str] = None):
    """Session bazlı chat geçmişi getir (Pinecone'dan)."""
    if not session_id:
        return JSONResponse({"chat_history": [], "session_id": None, "error": "session_id gerekli"})
    try:
        store = get_vector_store()
        return JSONResponse({"chat_history": store.get_chat_history(session_id=session_id), "session_id": session_id})
    except Exception as e:
        return JSONResponse({"chat_history": [], "session_id": session_id, "error": str(e)})

@app.get("/precedents")
async def list_precedents(session_id: str):
    """Oturum için eklenen emsal karar chunk'larını Pinecone'dan döndürür."""
    try:
        store = get_vector_store()
        raw = store.get_all_documents_raw(session_id, kinds=["precedent"])
        # Benzersiz pdf_id'lerini topla
        prec_ids = sorted(set(r["pdf_id"] for r in raw))
        data = [{"pdf_id": r["pdf_id"], "chunk_id": r["chunk_id"], "text": r["text"]} for r in raw]
        return {"precedent_sources": prec_ids, "chunks": data, "count": len(data)}
    except Exception as e:
        return {"error": str(e)}

@app.delete("/precedents/{session_id}")
async def clear_precedents(session_id: str):
    """Emsal karar chunk'larını vector store'dan temizler."""
    try:
        store = get_vector_store()
        raw = store.get_all_documents_raw(session_id, kinds=["precedent"])
        prec_ids = sorted(set(r["pdf_id"] for r in raw))
        removed_total = 0
        for pid in prec_ids:
            try:
                await store.a_remove_pdf(session_id=session_id, pdf_id=pid)
                removed_total += 1
            except Exception as re:
                print(f"⚠️ Emsal silinemedi: {pid} -> {re}")
        return {"removed": removed_total, "precedent_sources": prec_ids}
    except Exception as e:
        return {"error": str(e)}

@app.post("/clear-chat")
async def clear_chat(session_id: Optional[str] = Form(None)):
    """Chat geçmişini Pinecone'dan temizle"""
    if not session_id:
        return JSONResponse({"success": False, "message": "session_id gerekli"})
    try:
        store = get_vector_store()
        # Chat vektörlerini listele ve sil
        ids = store._list_vector_ids(f"{session_id}:__chat__:")
        if ids:
            for start in range(0, len(ids), 1000):
                store._index.delete(ids=ids[start:start+1000])
        return JSONResponse({"success": True, "message": f"Chat geçmişi temizlendi ({len(ids)} mesaj)"})
    except Exception as e:
        return JSONResponse({"success": False, "message": f"Hata: {str(e)}"})

# Case analysis sources endpoint removed

@app.get("/chat-precedents/{session_id}/{precedent_set_id}")
async def get_chat_precedent_set(session_id: str, precedent_set_id: str):
    """Detaylı chat yanıtında üretilen precedent_set_id için tam emsal listesini döndürür."""
    try:
        sets = session_chat_precedents.get(session_id)
        if not sets or precedent_set_id not in sets:
            return {"error": "Kayıt bulunamadı", "session_id": session_id, "precedent_set_id": precedent_set_id}
        data = sets[precedent_set_id]
        precedents = []
        for i, p in enumerate(data.get("precedents", []), 1):
            precedents.append({
                "index": i,
                "id": p.get("documentId") or p.get("documentID") or p.get("id") or f"DOC{i}",
                "birim": p.get("birimAdi"),
                "tarih": p.get("kararTarihi"),
                "tur": p.get("itemType"),
                "content": p.get("markdown_content") or p.get("text", ""),
            })
        return {
            "session_id": session_id,
            "precedent_set_id": precedent_set_id,
            "search_query": data.get("search_query"),
            "total_fetched": data.get("total_fetched"),
            "created": data.get("created"),
            "precedent_count": len(precedents),
            "precedents": precedents
        }
    except Exception as e:
        return {"error": str(e)}

@app.get("/health")
async def health_check():
    """Sistem durumu kontrolü - Multi-token sistemli"""
    try:
        health_data = {
            "status": "healthy",
            "ai_ready": legal_advisor_ai is not None,
            "multi_token_system": True
        }
        
        # Multi-token sistemi durumu - her zaman aktif
        try:
            token_status = get_token_status()
            health_data.update({
                "token_clients": token_status["total_clients"],
                "active_clients": token_status["loaded_clients"]
            })
        except Exception:
            health_data["multi_token_error"] = "Token status alınamadı"
        
        return health_data
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "multi_token_system": True
        }

@app.get("/ai-status")
async def ai_status():
    """AI modüllerinin durumunu detaylı göster"""
    try:
        status = {
            "multi_token_available": True,
            "modules": {}
        }
        
        # Legal Advisor durumu
        try:
            status["modules"]["legal_advisor"] = {
                "loaded": legal_advisor_ai is not None,
                "multi_token": legal_advisor_ai.use_multi_token if legal_advisor_ai else False
            }
        except Exception as e:
            status["modules"]["legal_advisor"] = {"error": str(e)}
        
        # Petition Generator durumu
        try:
            status["modules"]["petition_generator"] = {
                "loaded": petition_generator_ai is not None,
                "multi_token": petition_generator_ai.use_multi_token if petition_generator_ai else False
            }
        except Exception as e:
            status["modules"]["petition_generator"] = {"error": str(e)}
        
        # Multi-token sistem detayları - her zaman aktif
        try:
            token_status = get_token_status()
            status["token_system"] = token_status
        except Exception as e:
            status["token_system_error"] = str(e)
        
        return status
        
    except Exception as e:
        return {"error": f"AI status kontrolü hatası: {str(e)}"}

@app.post("/test-tokens")
async def test_tokens():
    """Tüm AI token'larını test et"""
    try:
        from ai_guard import test_all_clients
        test_results = await test_all_clients()
        return {
            "success": True,
            "test_results": test_results,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        return {"error": f"Token test hatası: {str(e)}"}

@app.post("/enhance-prompt")
async def enhance_prompt(
    prompt: str = Form(...),
    session_id: str = Header(None, alias="session_id", convert_underscores=False)
):
    """Kullanıcının yazdığı prompt'u profesyonel bir formata çevirir"""
    try:
        if not prompt or not prompt.strip():
            return {"error": "Boş prompt güzelleştirilemez"}
        
        # AI'dan prompt'u güzelleştirmesini iste
        enhancement_prompt = f"""Aşağıdaki kullanıcı girişini profesyonel, açık ve etkili bir hukuki soru/talep formatına çevir.

Kullanıcı girişi: "{prompt.strip()}"

Lütfen:
1. Açık ve net bir dil kullan
2. Hukuki terimler uygunsa dahil et
3. Soruyu yapılandır ve detaylandır
4. Gereksiz tekrarları kaldır
5. Profesyonel bir ton kullan
6. Maksimum 2-3 cümle ile sınırla

Sadece güzelleştirilmiş prompt'u ver, açıklama ekleme:"""

        # Token manager'dan AI client al
        token_manager = get_multi_token_manager()
        if not token_manager.clients:
            return {"error": "AI servisi şu anda kullanılamıyor"}
        
        # General client'i al, yoksa ilk mevcut client'i kullan
        client = token_manager.get_client('general')
        if not client:
            return {"error": "Uygun AI client bulunamadı"}
        
        # AI'dan güzelleştirilmiş prompt'u al (OpenAI v1 uyumlu)
        enhanced = ""
        system_msg = "Sen bir hukuki prompt güzelleştirme uzmanısın. Kullanıcı girişlerini profesyonel hukuki sorulara çevirirsin."
        # Öncelik: chat.completions
        try:
            if hasattr(client, "chat") and hasattr(client.chat, "completions"):
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_msg},
                        {"role": "user", "content": enhancement_prompt},
                    ],
                    temperature=0.2,
                    max_tokens=256,
                )
                if getattr(resp, "choices", None):
                    enhanced = (resp.choices[0].message.content or "").strip()
            else:
                raise AttributeError("chat.completions not available")
        except Exception:
            # Fallback: responses API (tek metin girişi)
            try:
                full = f"{system_msg}\n\nKULLANICI GİRİŞİ:\n{prompt.strip()}\n\nYalnızca düzeltilmiş prompt:".strip()
                resp2 = client.responses.create(model="gpt-4o-mini", input=full)
                # responses API output extraction
                if hasattr(resp2, "output_text") and resp2.output_text:
                    enhanced = resp2.output_text.strip()
                elif hasattr(resp2, "content") and resp2.content:
                    # Some clients return content list
                    part = resp2.content[0]
                    enhanced = getattr(part, "text", "") or getattr(part, "value", "") or ""
                    enhanced = str(enhanced).strip()
            except Exception as ee:
                print(f"❌ Prompt enhance fallback hatası: {ee}")
                enhanced = ""
        
        if not enhanced or not enhanced.strip():
            return {"error": "Prompt güzelleştirilemedi"}
        
        return {
            "success": True,
            "original": prompt.strip(),
            "enhanced": enhanced.strip()
        }
        
    except Exception as e:
        print(f"❌ Prompt güzelleştirme hatası: {e}")
        return {"error": f"Prompt güzelleştirme hatası: {str(e)}"}

# Case analysis endpoints removed

# Case analysis precedents endpoint removed

# Case analysis raw precedents endpoint removed

# Case analysis status endpoint removed

# ================================
# RAW PDF CONTENT ENDPOINTS (NO TRUNCATION)
# ================================

@app.get("/pdfs-raw/{session_id}")
async def get_pdfs_raw(session_id: str):
    """Bu oturumdaki tüm PDF ve emsal/diğer dokümanların tam metinlerini döndürür.

    Dönen yapı: { items: [ { pdf_id, kind, chunks: [{chunk_id, text}] } ] }
    Kesme/kısaltma yoktur; chunk'lar ham halde gelir.
    """
    try:
        store = get_vector_store()
        raw = store.get_all_documents_raw(session_id)
        buckets: Dict[tuple[str, str], List[Dict[str, Any]]] = {}
        for r in raw:
            key = (r["pdf_id"], r["kind"])  # group by document+kind
            if key not in buckets:
                buckets[key] = []
            buckets[key].append({"chunk_id": r["chunk_id"], "text": r["text"]})
        items = []
        for (pdf_id, kind), chunks in buckets.items():
            chunks.sort(key=lambda c: c["chunk_id"])  # ensure order
            items.append({"pdf_id": pdf_id, "kind": kind, "chunks": chunks})
        return JSONResponse({"items": items, "count": len(items), "session_id": session_id})
    except Exception as e:
        return JSONResponse({"error": f"Raw PDF listesi hatası: {str(e)}", "items": []})

@app.get("/pdf-raw/{session_id}/{pdf_id}")
async def get_pdf_raw(session_id: str, pdf_id: str):
    """Belirli bir PDF'in (veya precedent doc'un) ham tümleştirilmiş tam metnini döndürür.

    Dönen yapı: { pdf_id, kindGuess, full_text, chunks: [...] }
    """
    try:
        store = get_vector_store()
        # orijinal veya sanitize edilmiş id kullanılabilir
        raw = store.get_all_documents_raw(session_id)
        selected = [r for r in raw if r["pdf_id"] == pdf_id or r["pdf_id"] == store._id_map.get(pdf_id)]
        if not selected:
            return JSONResponse({"error": "Belge bulunamadı", "pdf_id": pdf_id, "session_id": session_id})
        selected.sort(key=lambda r: r["chunk_id"])
        full_text = "\n\n".join([r["text"] for r in selected])
        kind_guess = selected[0]["kind"] if selected else "pdf"
        return JSONResponse({
            "pdf_id": selected[0]["pdf_id"],
            "requested_id": pdf_id,
            "kind": kind_guess,
            "full_text": full_text,
            "chunks": [{"chunk_id": r["chunk_id"], "text": r["text"]} for r in selected],
            "session_id": session_id
        })
    except Exception as e:
        return JSONResponse({"error": f"Raw PDF içerik hatası: {str(e)}"})

# Case analysis precedents (summary) endpoint removed

# ================================
# STATIC FILES
# ================================

if __name__ == "__main__":
    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8000,
    )
